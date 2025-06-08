from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, Query
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from typing import List, Optional, Dict, Any
import json
import uuid
import aiofiles
import asyncio
from pathlib import Path
import PyPDF2
import docx
from datetime import datetime
import logging

from ..main import get_db, get_current_active_user, require_role, redis_client
from ..models import User, Document, AuditLog
from ..schemas import DocumentCreate, DocumentResponse, DocumentSearch

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/documents", tags=["documents"])

# Поддерживаемые типы файлов
ALLOWED_EXTENSIONS = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "txt": "text/plain",
    "md": "text/markdown"
}

DOCUMENT_CATEGORIES = [
    "ISO", "ГОСТ", "ФЗ", "Приказы", "NIST", "PCI DSS", 
    "GDPR", "OWASP", "CIS", "HIPAA", "COBIT"
]

def extract_text_from_file(file_path: Path, file_type: str) -> str:
    """Извлечение текста из различных типов файлов"""
    try:
        if file_type == "pdf":
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text()
                return text
        elif file_type == "docx":
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        elif file_type in ["txt", "md"]:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        return ""
    except Exception as e:
        logger.error(f"Ошибка извлечения текста из {file_path}: {e}")
        return ""

async def index_document_elasticsearch(document: Document, content: str):
    """Индексация документа в Elasticsearch"""
    try:
        # В реальном проекте здесь была бы интеграция с Elasticsearch
        # Пока используем Redis для кэширования поискового индекса
        if redis_client:
            search_data = {
                "id": document.id,
                "title": document.title,
                "content": content[:1000],  # Первые 1000 символов
                "category": document.category,
                "tags": json.loads(document.tags) if document.tags else [],
                "file_type": document.file_type
            }
            await redis_client.hset(
                f"doc_search:{document.id}", 
                mapping=search_data
            )
            
            # Добавляем в поисковый индекс по категориям
            await redis_client.sadd(
                f"category:{document.category}", 
                document.id
            )
    except Exception as e:
        logger.error(f"Ошибка индексации документа {document.id}: {e}")

@router.get("/", response_model=List[DocumentResponse])
async def get_documents(
    skip: int = Query(0, ge=0),
    limit: int = Query(100, le=1000),
    category: Optional[str] = Query(None),
    search: Optional[str] = Query(None),
    file_type: Optional[str] = Query(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user)
):
    """Получение списка документов с фильтрацией и поиском"""
    
    query = db.query(Document)
    
    # Фильтрация по категории
    if category and category in DOCUMENT_CATEGORIES:
        query = query.filter(Document.category == category)
    
    # Фильтрация по типу файла
    if file_type and file_type in ALLOWED_EXTENSIONS:
        query = query.filter(Document.file_type == file_type)
    
    # Поиск по названию и тегам
    if search:
        search_term = f"%{search}%"
        query = query.filter(
            Document.title.ilike(search_term) |
            Document.tags.ilike(search_term) |
            Document.content_preview.ilike(search_term)
        )
    
    # Пагинация
    documents = query.offset(skip).limit(limit).all()
    
    # Логирование поиска
    db.add(AuditLog(
        user_id=current_user.id,
        action="SEARCH_DOCUMENTS",
        resource="documents",
        details=json.dumps({
            "search": search,
            "category": category,
            "file_type": file_type,
            "results_count": len(documents)
        })
    ))
    db.commit()
    
    return documents

@router.post("/upload", response_model=DocumentResponse)
async def upload_document(
    title: str = Form(...),
    category: str = Form(...),
    tags: str = Form("[]"),  # JSON array as string
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_role(["admin", "super_admin"]))
):
    """Загрузка нового документа"""
    
    # Проверка типа файла
    file_extension = file.filename.split('.')[-1].lower()
    if file_extension not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Неподдерживаемый тип файла. Разрешены: {list(ALLOWED_EXTENSIONS.keys())}"
        )
    
    # Проверка размера файла (50MB)
    if file.size > 50 * 1024 * 1024:
        raise HTTPException(
            status_code=400,
            detail="Размер файла не должен превышать 50MB"
        )
    
    # Проверка категории
    if category not in DOCUMENT_CATEGORIES:
        raise HTTPException(
            status_code=400,
            detail=f"Неподдерживаемая категория. Разрешены: {DOCUMENT_CATEGORIES}"
        )
    
    try:
        # Парсинг тегов
        tags_list = json.loads(tags)
        if not isinstance(tags_list, list):
            raise ValueError("Tags должны быть массивом")
    except (json.JSONDecodeError, ValueError):
        raise HTTPException(
            status_code=400,
            detail="Теги должны быть в формате JSON массива"
        )
    
    # Создание уникального имени файла
    file_id = str(uuid.uuid4())
    filename = f"{file_id}.{file_extension}"
    file_path = Path("uploads/documents") / filename
    
    # Сохранение файла
    try:
        async with aiofiles.open(file_path, 'wb') as f:
            content = await file.read()
            await f.write(content)
    except Exception as e:
        logger.error(f"Ошибка сохранения файла: {e}")
        raise HTTPException(status_code=500, detail="Ошибка сохранения файла")
    
    # Извлечение текста для предварительного просмотра
    content_text = extract_text_from_file(file_path, file_extension)
    content_preview = content_text[:500] + "..." if len(content_text) > 500 else content_text
    
    # Создание записи в БД
    document = Document(
        title=title,
        filename=file.filename,
        file_path=str(file_path),
        category=category,
        tags=json.dumps(tags_list),
        content_preview=content_preview,
        file_size=file.size,
        file_type=file_extension,
        uploaded_by=current_user.id
    )
    
    db.add(document)
    db.commit()
    db.refresh(document)
    
    # Индексация для поиска
    await index_document_elasticsearch(document, content_text)
    
    # Логирование
    db.add(AuditLog(
        user_id=current_user.id,
        action="UPLOAD_DOCUMENT",
        resource="documents",
        resource_id=str(document.id),
        details=json.dumps({
            "title": title,
            "category": category,
            "file_size": file.size,
            "file_type": file_extension
        })
    ))
    db.commit()
    
    logger.info(f"Документ загружен: {title} (ID: {document.id})")
    
    return document

@router.get("/{document_id}", response_model=DocumentResponse)
async def get_document(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user)
):
    """Получение информации о документе"""
    
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Документ не найден")
    
    # Логирование просмотра
    db.add(AuditLog(
        user_id=current_user.id,
        action="VIEW_DOCUMENT",
        resource="documents",
        resource_id=str(document_id)
    ))
    db.commit()
    
    return document

@router.get("/{document_id}/download")
async def download_document(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user)
):
    """Скачивание файла документа"""
    
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Документ не найден")
    
    file_path = Path(document.file_path)
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Файл не найден на диске")
    
    # Логирование скачивания
    db.add(AuditLog(
        user_id=current_user.id,
        action="DOWNLOAD_DOCUMENT",
        resource="documents",
        resource_id=str(document_id)
    ))
    db.commit()
    
    return FileResponse(
        path=str(file_path),
        filename=document.filename,
        media_type=ALLOWED_EXTENSIONS.get(document.file_type, "application/octet-stream")
    )

@router.delete("/{document_id}")
async def delete_document(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_role(["admin", "super_admin"]))
):
    """Удаление документа"""
    
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Документ не найден")
    
    # Удаление файла с диска
    try:
        file_path = Path(document.file_path)
        if file_path.exists():
            file_path.unlink()
    except Exception as e:
        logger.error(f"Ошибка удаления файла {document.file_path}: {e}")
    
    # Удаление из поискового индекса
    try:
        if redis_client:
            await redis_client.delete(f"doc_search:{document_id}")
            await redis_client.srem(f"category:{document.category}", document_id)
    except Exception as e:
        logger.error(f"Ошибка удаления из индекса: {e}")
    
    # Удаление из БД
    db.delete(document)
    
    # Логирование
    db.add(AuditLog(
        user_id=current_user.id,
        action="DELETE_DOCUMENT",
        resource="documents",
        resource_id=str(document_id),
        details=json.dumps({"title": document.title})
    ))
    db.commit()
    
    return {"message": "Документ успешно удален"}

@router.get("/search/advanced")
async def advanced_search(
    query: str = Query(..., min_length=2),
    categories: Optional[List[str]] = Query(None),
    tags: Optional[List[str]] = Query(None),
    date_from: Optional[datetime] = Query(None),
    date_to: Optional[datetime] = Query(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user)
):
    """Расширенный поиск документов"""
    
    # Базовый поиск в БД
    db_query = db.query(Document)
    
    # Полнотекстовый поиск
    search_term = f"%{query}%"
    db_query = db_query.filter(
        Document.title.ilike(search_term) |
        Document.content_preview.ilike(search_term) |
        Document.tags.ilike(search_term)
    )
    
    # Фильтры
    if categories:
        db_query = db_query.filter(Document.category.in_(categories))
    
    if date_from:
        db_query = db_query.filter(Document.created_at >= date_from)
    
    if date_to:
        db_query = db_query.filter(Document.created_at <= date_to)
    
    results = db_query.limit(100).all()
    
    # Дополнительная фильтрация по тегам
    if tags:
        filtered_results = []
        for doc in results:
            doc_tags = json.loads(doc.tags) if doc.tags else []
            if any(tag in doc_tags for tag in tags):
                filtered_results.append(doc)
        results = filtered_results
    
    # Логирование
    db.add(AuditLog(
        user_id=current_user.id,
        action="ADVANCED_SEARCH",
        resource="documents",
        details=json.dumps({
            "query": query,
            "categories": categories,
            "tags": tags,
            "results_count": len(results)
        })
    ))
    db.commit()
    
    return {
        "query": query,
        "results": results,
        "total": len(results)
    }

@router.get("/categories/list")
async def get_categories(
    current_user: User = Depends(get_current_active_user)
):
    """Получение списка доступных категорий"""
    return {"categories": DOCUMENT_CATEGORIES}

@router.get("/stats/summary")
async def get_documents_stats(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user)
):
    """Статистика по документам"""
    
    total_docs = db.query(Document).count()
    
    # Статистика по категориям
    category_stats = {}
    for category in DOCUMENT_CATEGORIES:
        count = db.query(Document).filter(Document.category == category).count()
        if count > 0:
            category_stats[category] = count
    
    # Статистика по типам файлов
    file_type_stats = {}
    for file_type in ALLOWED_EXTENSIONS.keys():
        count = db.query(Document).filter(Document.file_type == file_type).count()
        if count > 0:
            file_type_stats[file_type] = count
    
    return {
        "total_documents": total_docs,
        "by_category": category_stats,
        "by_file_type": file_type_stats,
        "total_size_mb": sum(doc.file_size for doc in db.query(Document).all()) / 1024 / 1024
    } 